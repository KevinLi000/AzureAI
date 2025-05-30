"""
增强表格检测和文本对齐修复模块
用于解决PDF转Word中的表格识别错误、文字错位和样式丢失问题
"""

import os
import sys
import traceback
import types
import re
import fitz  # PyMuPDF

def apply_advanced_table_fixes(converter):
    """
    应用高级表格修复到PDF转换器
    
    参数:
        converter: PDF转换器实例
    
    返回:
        修改后的转换器实例
    """
    try:
        print("应用高级表格修复...")
        
        # 1. 增强表格文本对齐方法
        enhance_table_text_alignment(converter)
        
        # 2. 增强表格单元格合并方法
        enhance_table_cell_merging(converter)
        
        # 3. 增强表格结构检测
        enhance_table_structure_detection(converter)
        
        # 4. 增强表格文本提取
        enhance_table_text_extraction(converter)
        
        print("高级表格修复应用完成")
        return converter
        
    except Exception as e:
        print(f"应用高级表格修复时出错: {e}")
        traceback.print_exc()
        return converter

def enhance_table_text_alignment(converter):
    """增强表格文本对齐处理"""
    try:
        # 检查并增强表格单元格处理方法
        if hasattr(converter, '_process_table_block'):
            original_process_table = converter._process_table_block
              def enhanced_process_table_with_alignment(self, doc, block, page, pdf_document):
                """增强的表格处理方法，解决文本对齐问题"""
                try:
                    # 获取表格数据
                    table_data = block.get("table_data", [])
                    merged_cells = block.get("merged_cells", [])
                    
                    # 增强: 预处理表格数据以改进文本对齐
                    table_data, merged_cells = preprocess_table_data_for_alignment(table_data, merged_cells)
                    
                    # 更新表格数据
                    block["table_data"] = table_data
                    block["merged_cells"] = merged_cells
                    
                    # 检查表格处理方法中是否有特殊的换行符处理
                    if not hasattr(self, '_handle_table_newlines'):
                        # 添加处理换行符的方法
                        def handle_table_newlines(self, cell, text):
                            """处理表格单元格中的换行符"""
                            if '\n' not in text:
                                cell.text = text
                                return
                                
                            # 有换行符，需要特殊处理
                            # 清除单元格中的任何现有文本
                            for paragraph in cell.paragraphs:
                                if paragraph.text:
                                    paragraph._element.clear_content()
                            
                            # 分割文本并添加为多个段落
                            text_lines = text.split('\n')
                            for i, line in enumerate(text_lines):
                                if i == 0:
                                    # 使用第一个段落
                                    if cell.paragraphs:
                                        p = cell.paragraphs[0]
                                        p.text = line.strip()
                                    else:
                                        p = cell.add_paragraph(line.strip())
                                else:
                                    # 添加新段落
                                    p = cell.add_paragraph(line.strip())
                                
                                # 设置段落属性
                                from docx.enum.text import WD_ALIGN_PARAGRAPH
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                p.style = 'Normal'
                                
                        # 添加方法到转换器
                        import types
                        self._handle_table_newlines = types.MethodType(handle_table_newlines, self)
                    
                    # 调用原始方法
                    return original_process_table(self, doc, block, page, pdf_document)
                    
                except Exception as e:
                    print(f"增强表格文本对齐处理出错: {e}")
                    traceback.print_exc()
                    # 出错时调用原始方法
                    return original_process_table(self, doc, block, page, pdf_document)
            
            # 替换方法
            converter._process_table_block = types.MethodType(enhanced_process_table_with_alignment, converter)
            print("已应用表格文本对齐增强")
            
    except Exception as e:
        print(f"增强表格文本对齐处理失败: {e}")
        traceback.print_exc()

def enhance_table_cell_merging(converter):
    """增强表格单元格合并处理"""
    try:
        # 检查并增强表格单元格合并方法
        if hasattr(converter, '_detect_merged_cells'):
            original_detect_merged_cells = converter._detect_merged_cells
            
            def enhanced_detect_merged_cells(self, table):
                """增强的单元格合并检测方法"""
                try:
                    # 先使用原始方法
                    merged_cells = original_detect_merged_cells(self, table)
                    
                    # 增强: 修复可能的合并单元格问题
                    fixed_merged_cells = fix_merged_cells_issues(merged_cells, table)
                    
                    # 如果修复后的结果存在，使用修复后的结果
                    if fixed_merged_cells is not None:
                        return fixed_merged_cells
                    
                    return merged_cells
                    
                except Exception as e:
                    print(f"增强单元格合并检测出错: {e}")
                    traceback.print_exc()
                    # 出错时尝试使用原始方法
                    try:
                        return original_detect_merged_cells(self, table)
                    except:
                        return []
            
            # 替换方法
            converter._detect_merged_cells = types.MethodType(enhanced_detect_merged_cells, converter)
            print("已应用表格单元格合并增强")
            
    except Exception as e:
        print(f"增强表格单元格合并处理失败: {e}")
        traceback.print_exc()

def enhance_table_structure_detection(converter):
    """增强表格结构检测"""
    try:
        # 检查是否有_extract_tables方法
        if hasattr(converter, '_extract_tables'):
            original_extract_tables = converter._extract_tables
            
            def enhanced_extract_tables(self, pdf_document, page_num):
                """增强的表格提取方法"""
                try:
                    # 先使用原始方法
                    tables = original_extract_tables(self, pdf_document, page_num)
                    
                    # 如果没有检测到表格，尝试使用其他方法
                    if not tables or len(tables) == 0:
                        # 尝试使用增强型表格检测
                        enhanced_tables = extract_tables_advanced(self, pdf_document, page_num)
                        if enhanced_tables and len(enhanced_tables) > 0:
                            print(f"使用增强型表格检测成功识别表格")
                            return enhanced_tables
                    
                    return tables
                    
                except Exception as e:
                    print(f"增强表格提取出错: {e}")
                    traceback.print_exc()
                    # 出错时尝试使用替代方法
                    try:
                        return extract_tables_fallback(self, pdf_document, page_num)
                    except:
                        return []
            
            # 替换方法
            converter._extract_tables = types.MethodType(enhanced_extract_tables, converter)
            print("已应用表格结构检测增强")
            
    except Exception as e:
        print(f"增强表格结构检测失败: {e}")
        traceback.print_exc()

def enhance_table_text_extraction(converter):
    """增强表格文本提取"""
    try:
        # 添加表格文本提取增强方法
        def extract_table_text_enhanced(self, table, page):
            """
            增强型表格文本提取，确保文本在单元格内对齐正确
            
            参数:
                table: 表格对象
                page: 页面对象
                
            返回:
                提取的文本数据
            """
            try:
                # 获取表格区域
                if isinstance(table, dict) and "bbox" in table:
                    table_rect = table["bbox"]
                elif hasattr(table, 'bbox'):
                    table_rect = table.bbox
                elif hasattr(table, 'rect'):
                    table_rect = table.rect
                else:
                    return None
                
                # 从页面中提取区域内的文本
                import fitz
                rect = fitz.Rect(table_rect)
                
                # 使用不同方式提取文本并比较结果
                # 1. 使用DICT模式，获取详细的文本位置
                text_dict = page.get_text("dict", clip=rect)
                
                # 2. 使用BLOCKS模式，获取块级文本
                text_blocks = page.get_text("blocks", clip=rect)
                
                # 3. 使用TEXT模式，获取纯文本
                text_plain = page.get_text("text", clip=rect)
                
                # 处理提取的文本，构建表格数据
                # 使用get_text("dict")结果来获取精确的文本位置信息
                cells = []
                
                if text_dict and "blocks" in text_dict:
                    for block in text_dict["blocks"]:
                        if block["type"] == 0:  # 文本块
                            for line in block["lines"]:
                                text = ""
                                # 合并所有文本span
                                for span in line["spans"]:
                                    text += span["text"]
                                
                                # 创建单元格信息
                                if text.strip():
                                    cells.append({
                                        "bbox": (line["bbox"][0], line["bbox"][1], line["bbox"][2], line["bbox"][3]),
                                        "text": text.strip()
                                    })
                
                # 如果无法提取到细节文本，使用备用方法
                if not cells and text_blocks:
                    for block in text_blocks:
                        # block格式为 (x0, y0, x1, y1, "text", block_no, block_type)
                        if block[4].strip() and block[6] == 0:  # 确保有文本且是文本块
                            cells.append({
                                "bbox": (block[0], block[1], block[2], block[3]),
                                "text": block[4].strip()
                            })
                
                # 如果仍无法提取，使用最基础的文本
                if not cells and text_plain:
                    # 将纯文本分割成行
                    lines = text_plain.split('\n')
                    y_pos = rect[1]
                    line_height = (rect[3] - rect[1]) / max(len(lines), 1)
                    
                    for line in lines:
                        if line.strip():
                            cells.append({
                                "bbox": (rect[0], y_pos, rect[2], y_pos + line_height),
                                "text": line.strip()
                            })
                            y_pos += line_height
                
                return cells
                
            except Exception as e:
                print(f"增强表格文本提取出错: {e}")
                traceback.print_exc()
                return None
        
        # 添加方法到转换器
        converter.extract_table_text_enhanced = types.MethodType(extract_table_text_enhanced, converter)
        print("已应用表格文本提取增强")
        
    except Exception as e:
        print(f"增强表格文本提取失败: {e}")
        traceback.print_exc()

# 辅助函数

def preprocess_table_data_for_alignment(table_data, merged_cells):
    """
    预处理表格数据以改进文本对齐
    
    参数:
        table_data: 表格数据
        merged_cells: 合并单元格信息
        
    返回:
        处理后的表格数据和合并单元格信息
    """
    if not table_data:
        return [], []
    
    # 处理每个单元格的文本
    for i, row in enumerate(table_data):
        for j, cell_content in enumerate(row):
            if not isinstance(cell_content, str):
                # 确保内容是字符串
                table_data[i][j] = str(cell_content) if cell_content is not None else ""
                continue
                  # 处理单元格文本，改进对齐
            content = cell_content
            
            # 替换连续空格为单个空格，但保留换行符
            content = re.sub(r' {2,}', ' ', content)
            
            # 处理文本中的换行符，确保正确识别和保留
            # 先统一所有可能的换行符形式
            content = content.replace('\r\n', '\n').replace('\r', '\n')
            
            # 删除前后空白
            content = content.strip()
            
            # 对表格头部（第一行）进行居中处理标记
            if i == 0:
                # 在表格头部添加居中标记（这会在转换成Word时解析）
                content = content.strip()
                
            # 更新处理后的内容
            table_data[i][j] = content
    
    # 检查合并单元格是否有重叠，修复可能的问题
    if merged_cells:
        fixed_merged_cells = []
        cell_usage = {}  # 跟踪单元格使用情况
        
        for mc in sorted(merged_cells, key=lambda x: (x[0], x[1])):
            start_row, start_col, end_row, end_col = mc
            
            # 检查是否与之前的合并单元格重叠
            valid = True
            for r in range(start_row, end_row + 1):
                for c in range(start_col, end_col + 1):
                    cell_key = (r, c)
                    if cell_key in cell_usage:
                        valid = False
                        break
                if not valid:
                    break
            
            # 如果有效，添加到修复后的列表并标记单元格使用情况
            if valid:
                fixed_merged_cells.append(mc)
                for r in range(start_row, end_row + 1):
                    for c in range(start_col, end_col + 1):
                        cell_usage[(r, c)] = True
        
        merged_cells = fixed_merged_cells
    
    return table_data, merged_cells

def fix_merged_cells_issues(merged_cells, table):
    """
    修复合并单元格的问题
    
    参数:
        merged_cells: 原始合并单元格信息
        table: 表格对象
        
    返回:
        修复后的合并单元格信息
    """
    if not merged_cells:
        return None
    
    # 获取表格的行列数
    rows = 0
    cols = 0
    
    try:
        # 尝试获取表格尺寸
        if isinstance(table, dict) and "table_data" in table:
            rows = len(table["table_data"])
            cols = len(table["table_data"][0]) if rows > 0 else 0
        elif hasattr(table, 'extract') and callable(table.extract):
            table_data = table.extract()
            rows = len(table_data)
            cols = len(table_data[0]) if rows > 0 else 0
        elif hasattr(table, 'cells') and table.cells:
            # 查找最大的行和列索引
            max_row = max_col = 0
            for cell in table.cells:
                if hasattr(cell, 'span') and len(cell.span) >= 2:
                    max_row = max(max_row, cell.span[0])
                    max_col = max(max_col, cell.span[1])
            rows = max_row + 1
            cols = max_col + 1
    except Exception as e:
        print(f"获取表格尺寸出错: {e}")
        return None
    
    # 如果无法确定表格尺寸，返回原始数据
    if rows == 0 or cols == 0:
        return None
    
    # 验证并修复合并单元格
    fixed_merged_cells = []
    for merge_info in merged_cells:
        if not isinstance(merge_info, (list, tuple)) or len(merge_info) != 4:
            continue
            
        start_row, start_col, end_row, end_col = merge_info
        
        # 确保索引在有效范围内
        if (0 <= start_row <= end_row < rows and
            0 <= start_col <= end_col < cols):
            fixed_merged_cells.append((start_row, start_col, end_row, end_col))
    
    # 检查修复后的合并单元格是否有变化
    if len(fixed_merged_cells) != len(merged_cells):
        return fixed_merged_cells
    
    # 如果没有变化，返回None表示不需要修改
    return None

def extract_tables_advanced(converter, pdf_document, page_num):
    """
    使用高级方法提取表格
    
    参数:
        converter: 转换器实例
        pdf_document: PDF文档
        page_num: 页码
        
    返回:
        提取的表格列表
    """
    try:
        page = pdf_document[page_num]
        tables = []
        
        # 尝试使用OpenCV方法检测表格
        try:
            import cv2
            import numpy as np
            
            # 渲染页面为图像
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            
            # 转换为灰度图像
            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY) if img.shape[2] >= 3 else img
            
            # 二值化
            thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 11, 8)
            
            # 形态学操作，连接表格线
            kernel = np.ones((3, 3), np.uint8)
            dilated = cv2.dilate(thresh, kernel, iterations=1)
            
            # 查找轮廓
            contours, _ = cv2.findContours(dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # 处理检测到的轮廓
            page_width, page_height = page.rect.width, page.rect.height
            for contour in contours:
                area = cv2.contourArea(contour)
                if area > 1000:  # 过滤小轮廓
                    x, y, w, h = cv2.boundingRect(contour)
                    
                    # 转换坐标到PDF坐标系
                    x0 = x * page_width / pix.width
                    y0 = y * page_height / pix.height
                    x1 = (x + w) * page_width / pix.width
                    y1 = (y + h) * page_height / pix.height
                    
                    # 创建表格对象
                    table = {
                        "bbox": (x0, y0, x1, y1),
                        "cells": converter.extract_table_text_enhanced({"bbox": (x0, y0, x1, y1)}, page)
                    }
                    
                    # 只有在成功提取单元格时添加表格
                    if table["cells"] and len(table["cells"]) > 0:
                        tables.append(table)
        except Exception as e:
            print(f"OpenCV表格检测出错: {e}")
        
        # 如果OpenCV方法未能检测到表格，尝试文本分析方法
        if not tables:
            try:
                # 获取页面文本块
                blocks = page.get_text("dict")["blocks"]
                
                # 查找可能的表格区域（文本块的集合）
                text_lines = []
                for block in blocks:
                    if block["type"] == 0:  # 文本块
                        for line in block["lines"]:
                            line_bbox = line["bbox"]
                            text = "".join([span["text"] for span in line["spans"]])
                            text_lines.append((line_bbox, text))
                
                # 如果有足够的文本行，分析它们的位置关系
                if len(text_lines) > 5:
                    # 按y坐标排序
                    text_lines.sort(key=lambda x: x[0][1])
                    
                    # 查找行对齐的文本组（可能是表格）
                    row_groups = []
                    current_group = [text_lines[0]]
                    current_y = text_lines[0][0][1]
                    
                    for i in range(1, len(text_lines)):
                        line_bbox, _ = text_lines[i]
                        # 如果y坐标接近，认为是同一行
                        if abs(line_bbox[1] - current_y) < 5:
                            current_group.append(text_lines[i])
                        else:
                            # 开始新行
                            if len(current_group) > 1:
                                row_groups.append(current_group)
                            current_group = [text_lines[i]]
                            current_y = line_bbox[1]
                    
                    # 添加最后一组
                    if len(current_group) > 1:
                        row_groups.append(current_group)
                    
                    # 如果有多个行对齐的组，可能是表格
                    if len(row_groups) > 2:
                        # 计算表格边界
                        min_x = min([min([line[0][0] for line in group]) for group in row_groups])
                        min_y = min([min([line[0][1] for line in group]) for group in row_groups])
                        max_x = max([max([line[0][2] for line in group]) for group in row_groups])
                        max_y = max([max([line[0][3] for line in group]) for group in row_groups])
                        
                        # 创建表格
                        table = {
                            "bbox": (min_x, min_y, max_x, max_y),
                            "cells": []
                        }
                        
                        # 提取单元格
                        for group in row_groups:
                            for line_bbox, text in group:
                                table["cells"].append({
                                    "bbox": line_bbox,
                                    "text": text
                                })
                        
                        tables.append(table)
            except Exception as e:
                print(f"文本分析表格检测出错: {e}")
        
        # 返回结果
        if tables:
            class TableCollection:
                def __init__(self, tables_list):
                    self.tables = tables_list
            
            return TableCollection(tables)
        
        return []
        
    except Exception as e:
        print(f"高级表格提取出错: {e}")
        traceback.print_exc()
        return []

def extract_tables_fallback(converter, pdf_document, page_num):
    """
    备用的表格提取方法
    
    参数:
        converter: 转换器实例
        pdf_document: PDF文档
        page_num: 页码
        
    返回:
        提取的表格列表
    """
    try:
        import fitz
        page = pdf_document[page_num]
        
        # 获取页面文本块
        blocks = page.get_text("dict")["blocks"]
        
        # 查找可能的表格特征
        tables = []
        for block in blocks:
            # 检查是否包含多个有序的文本行
            if block["type"] == 0 and len(block["lines"]) > 3:
                lines = block["lines"]
                
                # 计算平均行高和行距
                line_heights = []
                line_spaces = []
                
                for i in range(len(lines) - 1):
                    line_heights.append(lines[i]["bbox"][3] - lines[i]["bbox"][1])
                    line_spaces.append(lines[i+1]["bbox"][1] - lines[i]["bbox"][3])
                
                if line_heights and line_spaces:
                    avg_height = sum(line_heights) / len(line_heights)
                    avg_space = sum(line_spaces) / len(line_spaces)
                    
                    # 如果行距接近一致，可能是表格
                    if max(line_spaces) - min(line_spaces) < 5 and avg_space < avg_height * 0.5:
                        # 创建表格对象
                        table = {
                            "bbox": block["bbox"],
                            "cells": []
                        }
                        
                        # 提取单元格
                        for line in lines:
                            text = "".join([span["text"] for span in line["spans"]])
                            table["cells"].append({
                                "bbox": line["bbox"],
                                "text": text
                            })
                        
                        tables.append(table)
        
        # 返回结果
        if tables:
            class TableCollection:
                def __init__(self, tables_list):
                    self.tables = tables_list
            
            return TableCollection(tables)
        
        return []
        
    except Exception as e:
        print(f"备用表格提取出错: {e}")
        traceback.print_exc()
        return []

# 测试方法 - 仅用于调试
def test_advanced_table_fixes(pdf_path):
    """测试高级表格修复"""
    try:
        import fitz
        from enhanced_pdf_converter import EnhancedPDFConverter
        
        # 创建转换器实例
        converter = EnhancedPDFConverter()
        
        # 应用高级表格修复
        converter = apply_advanced_table_fixes(converter)
        
        # 打开PDF文件
        pdf_document = fitz.open(pdf_path)
        
        # 处理第一页
        page_num = 0
        page = pdf_document[page_num]
        
        # 尝试提取表格
        tables = extract_tables_advanced(converter, pdf_document, page_num)
        
        # 打印结果
        if tables and hasattr(tables, 'tables'):
            print(f"检测到 {len(tables.tables)} 个表格")
            for i, table in enumerate(tables.tables):
                print(f"表格 {i+1}:")
                print(f"  位置: {table['bbox']}")
                print(f"  单元格数量: {len(table['cells'])}")
                for j, cell in enumerate(table['cells'][:5]):  # 只打印前5个单元格
                    print(f"    单元格 {j+1}: {cell['text']}")
                if len(table['cells']) > 5:
                    print(f"    ... 还有 {len(table['cells'])-5} 个单元格")
        else:
            print("未检测到表格")
        
        # 关闭PDF文件
        pdf_document.close()
        
    except Exception as e:
        print(f"测试高级表格修复出错: {e}")
        traceback.print_exc()

if __name__ == "__main__":
    # 如果作为独立脚本运行，执行测试
    if len(sys.argv) > 1:
        test_advanced_table_fixes(sys.argv[1])
    else:
        print("使用方法: python advanced_table_fixes.py <pdf_file>")
